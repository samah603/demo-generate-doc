from docx import Document
import os

# Load template
doc = Document("WI_VDOOTTHUB_Deployment_template.docx")

# Get values from environment
tag_version = os.getenv("GITHUB_REF_NAME", "vX.Y.Z")
commit_message = os.getenv("GITHUB_HEAD_COMMIT_MESSAGE", "Auto-generated commit")
repo_name = os.getenv("GITHUB_REPOSITORY", "unknown-repo").split("/")[-1]

def is_drawing_run(run):
    return bool(run._element.xpath(".//w:drawing | .//w:pict"))

def replace_in_runs(paragraph):
    for run in paragraph.runs:
        if is_drawing_run(run):
            continue  # Skip runs that contain images
        text = run.text
        if not text:
            continue
        text = text.replace("{{TAGVERSION}}", tag_version)
        text = text.replace("{{PRMESSAGE}}", commit_message)
        text = text.replace("{{REPONAME}}", repo_name)
        run.text = text

# main document paragraphs
for p in doc.paragraphs:
    replace_in_runs(p)

# tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_in_runs(p)

# headers/footers
for section in doc.sections:
    for p in section.header.paragraphs:
        replace_in_runs(p)
    for p in section.footer.paragraphs:
        replace_in_runs(p)

# Save output
output_filename = f"WI_VDOOTTHUB_Deployment_{tag_version}.docx"
doc.save(output_filename)
